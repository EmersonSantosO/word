import kivy
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from kivy.uix.filechooser import FileChooserListView
from kivy.uix.popup import Popup
from kivy.uix.scrollview import ScrollView
from kivy.uix.colorpicker import ColorPicker
from kivy.uix.togglebutton import ToggleButton
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK, WD_UNDERLINE
import os
import ctypes
from ctypes import wintypes

# Algun Comentario de Prueba para el repositorio de Github
# Otro Comentario de Prueba para el repositorio de Github 2
def is_file_accessible(file_path, mode="r"):
    if os.path.exists(file_path) is False:
        return False

    if os.access(file_path, os.R_OK) is False:
        return False
    try:
        f = open(file_path, mode)
        f.close()
    except IOError as e:
        return False
    return True


class WordCreatorApp(App):

    def build(self, *args, **kwargs):

        self.document = Document()
        self.title = "Word Creator App"
        self.selected_color = (0, 0, 0, 1)
        main_layout = BoxLayout(orientation="horizontal", padding=10, spacing=10)
        main_layout.add_widget(self.create_control_panel())
        main_layout.add_widget(self.create_preview_layout())
        return main_layout

    def create_control_panel(self):
        control_panel = BoxLayout(
            orientation="vertical", size_hint=(0.4, 1), padding=10, spacing=10
        )

        self.text_input = TextInput(
            hint_text="Ingrese su texto aquí", multiline=True, size_hint=(1, 0.3)
        )
        control_panel.add_widget(self.text_input)
        self.text_input.bind(text=self.update_preview)

        control_panel.add_widget(Label(text="Tamaño de Fuente:", size_hint=(1, 0.05)))
        self.font_size_spinner = Spinner(
            text="12", values=[str(i) for i in range(8, 41, 2)], size_hint=(1, 0.1)
        )
        self.font_size_spinner.bind(text=self.update_preview)
        control_panel.add_widget(self.font_size_spinner)

        control_panel.add_widget(Label(text="Estilo de Fuente:", size_hint=(1, 0.05)))
        self.font_style_spinner = Spinner(
            text="Normal",
            values=("Normal", "Negrita", "Cursiva", "Subrayado"),
            size_hint=(1, 0.1),
        )
        self.font_style_spinner.bind(text=self.update_preview)
        control_panel.add_widget(self.font_style_spinner)

        color_picker_button = Button(
            text="Seleccionar Color de Texto", size_hint=(1, 0.1)
        )
        color_picker_button.bind(on_press=self.open_color_picker)
        control_panel.add_widget(color_picker_button)

        control_panel.add_widget(Label(text="Alineación:", size_hint=(1, 0.05)))
        alignment_layout = BoxLayout(size_hint=(1, 0.1))
        self.align_left = ToggleButton(
            text="Izquierda", group="align", state="down", size_hint=(1, 1)
        )
        self.align_center = ToggleButton(text="Centro", group="align", size_hint=(1, 1))
        self.align_right = ToggleButton(text="Derecha", group="align", size_hint=(1, 1))
        self.align_left.bind(state=self.update_preview)
        self.align_center.bind(state=self.update_preview)
        self.align_right.bind(state=self.update_preview)
        alignment_layout.add_widget(self.align_left)
        alignment_layout.add_widget(self.align_center)
        alignment_layout.add_widget(self.align_right)
        control_panel.add_widget(alignment_layout)

        insert_image_button = Button(text="Insertar Imagen", size_hint=(1, 0.1))
        insert_image_button.bind(on_press=self.insert_image)
        control_panel.add_widget(insert_image_button)

        insert_table_button = Button(text="Insertar Tabla", size_hint=(1, 0.1))
        insert_table_button.bind(on_press=self.insert_table)
        control_panel.add_widget(insert_table_button)

        insert_header_button = Button(text="Insertar Encabezado", size_hint=(1, 0.1))
        insert_header_button.bind(on_press=self.insert_header)
        control_panel.add_widget(insert_header_button)

        insert_footer_button = Button(text="Insertar Pie de Página", size_hint=(1, 0.1))
        insert_footer_button.bind(on_press=self.insert_footer)
        control_panel.add_widget(insert_footer_button)

        insert_bullet_list_button = Button(
            text="Insertar Lista con Viñetas", size_hint=(1, 0.1)
        )
        insert_bullet_list_button.bind(on_press=self.insert_bullet_list)
        control_panel.add_widget(insert_bullet_list_button)

        insert_numbered_list_button = Button(
            text="Insertar Lista Numerada", size_hint=(1, 0.1)
        )
        insert_numbered_list_button.bind(on_press=self.insert_numbered_list)
        control_panel.add_widget(insert_numbered_list_button)

        save_button = Button(text="Guardar en Word", size_hint=(1, 0.1))
        save_button.bind(on_press=self.save_to_word)
        control_panel.add_widget(save_button)

        return control_panel

    def create_preview_layout(self):
        preview_layout = BoxLayout(
            orientation="vertical", size_hint=(0.6, 1), padding=10, spacing=10
        )
        preview_layout.add_widget(Label(text="Vista Previa:", size_hint=(1, 0.05)))

        scroll_view = ScrollView(size_hint=(1, 1))
        self.preview_label = Label(
            text="", size_hint=(1, None), valign="top", halign="left", markup=True
        )
        self.preview_label.bind(
            texture_size=lambda *x: setattr(
                self.preview_label, "height", self.preview_label.texture_size[1]
            )
        )
        scroll_view.add_widget(self.preview_label)
        preview_layout.add_widget(scroll_view)

        return preview_layout

    def update_preview(self, *args, **kwargs):
        try:
            font_size = int(self.font_size_spinner.text)
            font_style = self.font_style_spinner.text.lower()
            color_hex = f"{int(self.selected_color[0] * 255):02x}{int(self.selected_color[1] * 255):02x}{int(self.selected_color[2] * 255):02x}"

            font_style_tags = {
                "normal": "",
                "negrita": "[b]",
                "cursiva": "[i]",
                "subrayado": "[u]",
            }
            estilo_actual = font_style_tags.get(font_style, "")

            # Actualizamos la vista previa del texto
            preview_text = f"[color={color_hex}][size={font_size}]{estilo_actual}{self.text_input.text}{estilo_actual and '[/' + estilo_actual[1:]}"
            self.preview_label.text = preview_text

            if self.align_center.state == "down":
                self.preview_label.halign = "center"
            elif self.align_right.state == "down":
                self.preview_label.halign = "right"
            else:
                self.preview_label.halign = "left"

        except ValueError as e:
            print(f"Error de formateo en la vista previa: {e}")

    ## aqui se crea la funcion para seleccionar el color
    def open_color_picker(self, instance, *args, **kwargs):
        self.color_picker = ColorPicker(color=self.selected_color, size_hint=(1, 1))
        popup = Popup(
            ## aqui se crea el popup para seleccionar el color
            title="Seleccionar Color de Texto",
            content=self.color_picker,
            size_hint=(0.9, 0.9),
        )

        self.color_picker.bind(
            color=lambda instance, value: self.on_color_select(
                popup, value, *args, **kwargs
            )
        )
        popup.open()

    def on_color_select(self, popup, value):
        self.selected_color = value
        popup.dismiss()
        self.update_preview()
        if self.align_center.state == "down":

            self.preview_label.halign = "center"
        elif self.align_right.state == "down":
            self.preview_label.halign = "right"
        else:
            self.preview_label.halign = "left"
            pass

    def insert_image(self, instance):
        filechooser = FileChooserListView(size_hint=(1, 0.8))
        popup = Popup(
            title="Seleccionar Imagen", content=filechooser, size_hint=(0.9, 0.9)
        )
        filechooser.bind(
            on_selection=lambda x: self.on_image_select(popup, x.selection)
        )
        popup.open()

    def on_image_select(self, popup, selection):
        if selection and is_file_accessible(selection[0], "rb"):
            self.document.add_picture(selection[0], width=Inches(1.25))
            popup.dismiss()

    def insert_table(self, instance):
        popup = Popup(title="Insertar Tabla", size_hint=(0.9, 0.4))
        layout = BoxLayout(orientation="vertical", padding=10, spacing=10)

        rows_input = TextInput(
            hint_text="Número de filas", multiline=False, size_hint=(1, 0.2)
        )
        columns_input = TextInput(
            hint_text="Número de columnas", multiline=False, size_hint=(1, 0.2)
        )
        insert_button = Button(text="Insertar", size_hint=(1, 0.2))
        insert_button.bind(
            on_press=lambda x: self.on_table_insert(
                popup, rows_input.text, columns_input.text
            )
        )

        layout.add_widget(rows_input)
        layout.add_widget(columns_input)
        layout.add_widget(insert_button)

        popup.content = layout
        popup.open()

    def on_table_insert(self, popup, rows, columns):
        try:
            rows, columns = int(rows), int(columns)
            table = self.document.add_table(rows=rows, cols=columns)
            for row in table.rows:
                for cell in row.cells:
                    cell.text = "Texto"
            popup.dismiss()
        except ValueError:
            print("Número de filas y columnas debe ser un entero.")

    ## aqui se crea la funcion para insertar el encabezado
    def insert_header(self, instance, *args, **kwargs):

        self.document.sections[0].header.paragraphs[0].text = "Encabezado"

    ## aqui se crea la funcion para insertar el pie de pagina
    def insert_footer(self, instance, *args, **kwargs):

        self.document.sections[0].footer.paragraphs[0].text = "Pie de Página"

    def insert_bullet_list(self, instance):
        popup = Popup(title="Insertar Lista con Viñetas", size_hint=(0.9, 0.4))
        layout = BoxLayout(orientation="vertical", padding=10, spacing=10)

        bullet_text = TextInput(
            hint_text="Texto de la viñeta", multiline=True, size_hint=(1, 0.7)
        )
        insert_button = Button(text="Insertar", size_hint=(1, 0.2))
        insert_button.bind(
            on_press=lambda x: self.on_bullet_list_insert(popup, bullet_text.text)
        )

        layout.add_widget(bullet_text)
        layout.add_widget(insert_button)

        popup.content = layout
        popup.open()

    def on_bullet_list_insert(self, popup, bullet_text):
        for line in bullet_text.split("\n"):
            self.document.add_paragraph(line, style="ListBullet")
        popup.dismiss()

    def insert_numbered_list(self, instance):
        popup = Popup(title="Insertar Lista Numerada", size_hint=(0.9, 0.4))
        layout = BoxLayout(orientation="vertical", padding=10, spacing=10)

        numbered_text = TextInput(
            hint_text="Texto de la lista", multiline=True, size_hint=(1, 0.7)
        )
        insert_button = Button(text="Insertar", size_hint=(1, 0.2))
        insert_button.bind(
            on_press=lambda x: self.on_numbered_list_insert(popup, numbered_text.text)
        )

        layout.add_widget(numbered_text)
        layout.add_widget(insert_button)

        popup.content = layout
        popup.open()

    def on_numbered_list_insert(self, popup, numbered_text):
        for line in numbered_text.split("\n"):
            self.document.add_paragraph(line, style="ListNumber")

        popup.dismiss()

    def save_to_word(self, instance):
        filechooser = FileChooserListView(size_hint=(1, 0.8), path=".")
        popup = Popup(
            title="Guardar Archivo", content=filechooser, size_hint=(0.9, 0.9)
        )
        filechooser.bind(on_selection=lambda x: self.on_file_save(popup, x.selection))
        popup.open()

    def on_file_save(self, popup, selection):
        if selection:
            try:
                self.document.save(selection[0])
                popup.dismiss()
            except Exception as e:
                print(f"Error al guardar el archivo: {e}")
                ctypes.windll.user32.MessageBoxW(
                    0,
                    "Error al guardar el archivo. Asegúrese de que el archivo no esté abierto.",
                    "Error",
                    0 | 0x30,
                )
                popup.dismiss()


if __name__ == "__main__":
    WordCreatorApp().run()
